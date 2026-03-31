import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 简历生成器主类
class ResumeGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("定制化简历生成器")
        self.root.geometry("1200x800")
        
        # 创建全局样式
        self.style = ttk.Style()
        self.style.configure("TLabel", font=("宋体", 10))
        self.style.configure("TButton", font=("宋体", 10))
        self.style.configure("TEntry", font=("宋体", 10))
        
        # 初始化数据存储字典
        self.resume_data = {
            # 基本信息
            "姓名": "", "出生年月": "", "邮箱": "", "电话": "", 
            "籍贯": "", "政治面貌": "", "求职意向": "",
            # 教育经历
            "本科学校": "", "本科专业": "", "本科学位": "", "本科时间": "",
            "硕士学校": "", "硕士专业": "", "硕士学位": "", "硕士时间": "",
            "核心课程": "",
            # 专业技能
            "专业技能": "",
            # 项目经历
            "项目名称": "", "项目时间": "", "项目描述": "",
            # 获奖情况
            "获奖情况": "",
            # 论文与专利
            "论文专利": "",
            # 等级考试
            "等级考试": ""
        }
        
        self._create_widgets()
    
    def _create_widgets(self):
        # 1. 基本信息区域
        frame_base = ttk.LabelFrame(self.root, text="基本信息", padding=10)
        frame_base.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")
        
        # 基本信息表单
        base_labels = ["姓名", "出生年月", "邮箱", "电话", "籍贯", "政治面貌", "求职意向"]
        for idx, label_text in enumerate(base_labels):
            ttk.Label(frame_base, text=label_text + ":").grid(row=idx//2, column=(idx%2)*2, padx=5, pady=5, sticky="e")
            entry = ttk.Entry(frame_base, width=30)
            entry.grid(row=idx//2, column=(idx%2)*2 + 1, padx=5, pady=5, sticky="w")
            # 绑定输入框到数据字典
            entry.bind("<KeyRelease>", lambda e, k=label_text: self._update_data(k, e.widget.get()))
            setattr(self, f"entry_{label_text}", entry)
        
        # 2. 教育经历区域
        frame_edu = ttk.LabelFrame(self.root, text="教育经历", padding=10)
        frame_edu.grid(row=1, column=0, padx=10, pady=10, sticky="nsew")
        
        # 本科信息
        ttk.Label(frame_edu, text="本科阶段:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        edu_under_labels = ["本科学校", "本科专业", "本科学位", "本科时间"]
        for idx, label_text in enumerate(edu_under_labels):
            ttk.Label(frame_edu, text=label_text + ":").grid(row=1, column=idx*2, padx=5, pady=5, sticky="e")
            entry = ttk.Entry(frame_edu, width=20)
            entry.grid(row=1, column=idx*2 + 1, padx=5, pady=5, sticky="w")
            entry.bind("<KeyRelease>", lambda e, k=label_text: self._update_data(k, e.widget.get()))
            setattr(self, f"entry_{label_text}", entry)
        
        # 硕士信息
        ttk.Label(frame_edu, text="硕士阶段:").grid(row=2, column=0, padx=5, pady=5, sticky="w")
        edu_master_labels = ["硕士学校", "硕士专业", "硕士学位", "硕士时间"]
        for idx, label_text in enumerate(edu_master_labels):
            ttk.Label(frame_edu, text=label_text + ":").grid(row=3, column=idx*2, padx=5, pady=5, sticky="e")
            entry = ttk.Entry(frame_edu, width=20)
            entry.grid(row=3, column=idx*2 + 1, padx=5, pady=5, sticky="w")
            entry.bind("<KeyRelease>", lambda e, k=label_text: self._update_data(k, e.widget.get()))
            setattr(self, f"entry_{label_text}", entry)
        
        # 核心课程
        ttk.Label(frame_edu, text="核心课程:").grid(row=4, column=0, padx=5, pady=5, sticky="e")
        entry_core_course = ttk.Entry(frame_edu, width=80)
        entry_core_course.grid(row=4, column=1, columnspan=7, padx=5, pady=5, sticky="w")
        entry_core_course.bind("<KeyRelease>", lambda e: self._update_data("核心课程", e.widget.get()))
        self.entry_核心课程 = entry_core_course
        
        # 3. 详细内容区域（右侧大区域）
        frame_detail = ttk.LabelFrame(self.root, text="详细内容", padding=10)
        frame_detail.grid(row=0, column=1, rowspan=3, padx=10, pady=10, sticky="nsew")
        
        # 标签页切换
        notebook = ttk.Notebook(frame_detail)
        notebook.pack(fill="both", expand=True, padx=5, pady=5)
        
        # 3.1 专业技能标签页
        frame_skill = ttk.Frame(notebook)
        notebook.add(frame_skill, text="专业技能")
        ttk.Label(frame_skill, text="专业技能（每行一条）:").pack(anchor="w", padx=5, pady=5)
        self.text_skill = scrolledtext.ScrolledText(frame_skill, width=60, height=10, font=("宋体", 10))
        self.text_skill.pack(fill="both", expand=True, padx=5, pady=5)
        self.text_skill.bind("<KeyRelease>", lambda e: self._update_data("专业技能", self.text_skill.get(1.0, tk.END)))
        
        # 3.2 项目经历标签页
        frame_project = ttk.Frame(notebook)
        notebook.add(frame_project, text="项目经历")
        ttk.Label(frame_project, text="项目名称:").grid(row=0, column=0, padx=5, pady=5, sticky="e")
        entry_project_name = ttk.Entry(frame_project, width=40)
        entry_project_name.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        entry_project_name.bind("<KeyRelease>", lambda e: self._update_data("项目名称", e.widget.get()))
        self.entry_项目名称 = entry_project_name
        
        ttk.Label(frame_project, text="项目时间:").grid(row=0, column=2, padx=5, pady=5, sticky="e")
        entry_project_time = ttk.Entry(frame_project, width=20)
        entry_project_time.grid(row=0, column=3, padx=5, pady=5, sticky="w")
        entry_project_time.bind("<KeyRelease>", lambda e: self._update_data("项目时间", e.widget.get()))
        self.entry_项目时间 = entry_project_time
        
        ttk.Label(frame_project, text="项目描述（背景+工作+结果）:").grid(row=1, column=0, padx=5, pady=5, sticky="nw")
        self.text_project = scrolledtext.ScrolledText(frame_project, width=60, height=8, font=("宋体", 10))
        self.text_project.grid(row=1, column=1, columnspan=3, padx=5, pady=5, sticky="nsew")
        self.text_project.bind("<KeyRelease>", lambda e: self._update_data("项目描述", self.text_project.get(1.0, tk.END)))
        
        # 3.3 获奖情况标签页
        frame_award = ttk.Frame(notebook)
        notebook.add(frame_award, text="获奖情况")
        ttk.Label(frame_award, text="获奖情况（每行一条，自动添加项目符号）:").pack(anchor="w", padx=5, pady=5)
        self.text_award = scrolledtext.ScrolledText(frame_award, width=60, height=10, font=("宋体", 10))
        self.text_award.pack(fill="both", expand=True, padx=5, pady=5)
        self.text_award.bind("<KeyRelease>", lambda e: self._update_data("获奖情况", self.text_award.get(1.0, tk.END)))
        
        # 3.4 论文专利标签页
        frame_paper = ttk.Frame(notebook)
        notebook.add(frame_paper, text="论文与专利")
        ttk.Label(frame_paper, text="论文/专利（每行一条）:").pack(anchor="w", padx=5, pady=5)
        self.text_paper = scrolledtext.ScrolledText(frame_paper, width=60, height=10, font=("宋体", 10))
        self.text_paper.pack(fill="both", expand=True, padx=5, pady=5)
        self.text_paper.bind("<KeyRelease>", lambda e: self._update_data("论文专利", self.text_paper.get(1.0, tk.END)))
        
        # 3.5 等级考试标签页
        frame_exam = ttk.Frame(notebook)
        notebook.add(frame_exam, text="等级考试")
        ttk.Label(frame_exam, text="等级考试（每行一条）:").pack(anchor="w", padx=5, pady=5)
        self.text_exam = scrolledtext.ScrolledText(frame_exam, width=60, height=10, font=("宋体", 10))
        self.text_exam.pack(fill="both", expand=True, padx=5, pady=5)
        self.text_exam.bind("<KeyRelease>", lambda e: self._update_data("等级考试", self.text_exam.get(1.0, tk.END)))
        
        # 4. 生成按钮区域
        frame_btn = ttk.Frame(self.root)
        frame_btn.grid(row=3, column=0, columnspan=2, padx=10, pady=10, sticky="nsew")
        btn_generate = ttk.Button(frame_btn, text="生成简历（Word格式）", command=self.generate_resume)
        btn_generate.pack(side="left", padx=20, pady=10)
        
        btn_clear = ttk.Button(frame_btn, text="清空所有内容", command=self.clear_all)
        btn_clear.pack(side="left", padx=20, pady=10)
    
    def _update_data(self, key, value):
        """更新简历数据字典"""
        self.resume_data[key] = value.strip()
    
    def clear_all(self):
        """清空所有输入框内容"""
        for key in self.resume_data.keys():
            if hasattr(self, f"entry_{key}"):
                getattr(self, f"entry_{key}").delete(0, tk.END)
        # 清空文本框
        self.text_skill.delete(1.0, tk.END)
        self.text_project.delete(1.0, tk.END)
        self.text_award.delete(1.0, tk.END)
        self.text_paper.delete(1.0, tk.END)
        self.text_exam.delete(1.0, tk.END)
        # 重置数据字典
        self.resume_data = {k: "" for k in self.resume_data.keys()}
        messagebox.showinfo("提示", "所有内容已清空！")
    
    def generate_resume(self):
        """生成Word格式的简历"""
        # 创建Word文档
        doc = Document()
        
        # 设置文档字体（解决中文乱码）
        doc.styles['Normal'].font.name = '宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        doc.styles['Normal'].font.size = Pt(10.5)
        
        # ========== 1. 基本信息 ==========
        # 姓名行
        para_name = doc.add_paragraph()
        para_name.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_name = para_name.add_run("姓名\t")
        run_name.font.bold = True
        para_name.add_run(self.resume_data["姓名"])
        
        # 出生年月/电话行
        para_base1 = doc.add_paragraph()
        run_birth = para_base1.add_run("出生年月: ")
        run_birth.font.bold = True
        para_base1.add_run(self.resume_data["出生年月"] + "\t\t")
        
        run_phone = para_base1.add_run("电话: ")
        run_phone.font.bold = True
        para_base1.add_run(self.resume_data["电话"])
        
        # 邮箱/籍贯政治面貌行
        para_base2 = doc.add_paragraph()
        run_email = para_base2.add_run("邮箱: ")
        run_email.font.bold = True
        para_base2.add_run(self.resume_data["邮箱"] + "\t\t")
        
        run_other = para_base2.add_run("籍贯、政治面貌: ")
        run_other.font.bold = True
        para_base2.add_run(f"{self.resume_data['籍贯']}、{self.resume_data['政治面貌']}")
        
        # 求职意向
        if self.resume_data["求职意向"]:
            para_intention = doc.add_paragraph()
            run_intention = para_intention.add_run("求职意向: ")
            run_intention.font.bold = True
            para_intention.add_run(self.resume_data["求职意向"])
        
        doc.add_paragraph("-"*80)  # 分隔线
        
        # ========== 2. 教育经历 ==========
        para_edu_title = doc.add_paragraph("教育经历")
        para_edu_title.runs[0].font.bold = True
        para_edu_title.runs[0].font.size = Pt(12)
        
        # 本科经历
        if self.resume_data["本科学校"]:
            para_under = doc.add_paragraph()
            para_under.add_run(f"{self.resume_data['本科学校']}\t").bold = True
            para_under.add_run(f"{self.resume_data['本科专业']}\t")
            para_under.add_run(f"{self.resume_data['本科学位']}\t")
            para_under.add_run(self.resume_data["本科时间"])
        
        # 硕士经历
        if self.resume_data["硕士学校"]:
            para_master = doc.add_paragraph()
            para_master.add_run(f"{self.resume_data['硕士学校']}\t").bold = True
            para_master.add_run(f"{self.resume_data['硕士专业']}\t")
            para_master.add_run(f"{self.resume_data['硕士学位']}\t")
            para_master.add_run(self.resume_data["硕士时间"])
        
        # 核心课程
        if self.resume_data["核心课程"]:
            para_core = doc.add_paragraph()
            para_core.add_run("核心课程: ").bold = True
            para_core.add_run(self.resume_data["核心课程"])
        
        doc.add_paragraph("-"*80)
        
        # ========== 3. 专业技能 ==========
        para_skill_title = doc.add_paragraph("专业技能")
        para_skill_title.runs[0].font.bold = True
        para_skill_title.runs[0].font.size = Pt(12)
        
        # 拆分每行技能并添加项目符号
        skills = self.resume_data["专业技能"].split("\n")
        for skill in skills:
            if skill.strip():
                para_skill = doc.add_paragraph(f"● {skill.strip()}", style="List Bullet")
                para_skill.paragraph_format.left_indent = Inches(0.1)
        
        doc.add_paragraph("-"*80)
        
        # ========== 4. 项目经历 ==========
        para_project_title = doc.add_paragraph("项目经历")
        para_project_title.runs[0].font.bold = True
        para_project_title.runs[0].font.size = Pt(12)
        
        if self.resume_data["项目名称"]:
            para_project = doc.add_paragraph()
            para_project.add_run(f"◆ {self.resume_data['项目名称']}\t").bold = True
            para_project.add_run(self.resume_data["项目时间"])
            
            # 项目描述
            project_desc = self.resume_data["项目描述"].split("\n")
            for desc in project_desc:
                if desc.strip():
                    para_desc = doc.add_paragraph(f"● {desc.strip()}", style="List Bullet")
                    para_desc.paragraph_format.left_indent = Inches(0.2)
        
        doc.add_paragraph("-"*80)
        
        # ========== 5. 获奖情况 ==========
        para_award_title = doc.add_paragraph("获奖情况")
        para_award_title.runs[0].font.bold = True
        para_award_title.runs[0].font.size = Pt(12)
        
        awards = self.resume_data["获奖情况"].split("\n")
        for award in awards:
            if award.strip():
                para_award = doc.add_paragraph(f"● {award.strip()}", style="List Bullet")
                para_award.paragraph_format.left_indent = Inches(0.1)
        
        doc.add_paragraph("-"*80)
        
        # ========== 6. 论文与专利 ==========
        para_paper_title = doc.add_paragraph("论文与专利")
        para_paper_title.runs[0].font.bold = True
        para_paper_title.runs[0].font.size = Pt(12)
        
        papers = self.resume_data["论文专利"].split("\n")
        for paper in papers:
            if paper.strip():
                para_paper = doc.add_paragraph(f"● {paper.strip()}", style="List Bullet")
                para_paper.paragraph_format.left_indent = Inches(0.1)
        
        doc.add_paragraph("-"*80)
        
        # ========== 7. 等级考试 ==========
        para_exam_title = doc.add_paragraph("等级考试")
        para_exam_title.runs[0].font.bold = True
        para_exam_title.runs[0].font.size = Pt(12)
        
        exams = self.resume_data["等级考试"].split("\n")
        for exam in exams:
            if exam.strip():
                para_exam = doc.add_paragraph(f"● {exam.strip()}", style="List Bullet")
                para_exam.paragraph_format.left_indent = Inches(0.1)
        
        # 保存文档
        try:
            save_path = "定制化简历.docx"
            doc.save(save_path)
            messagebox.showinfo("成功", f"简历已生成！保存路径：{save_path}")
        except Exception as e:
            messagebox.showerror("错误", f"生成失败：{str(e)}")

# 程序入口
if __name__ == "__main__":
    root = tk.Tk()
    app = ResumeGenerator(root)
    root.mainloop()
